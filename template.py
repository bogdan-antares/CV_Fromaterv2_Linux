from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_TAB_LEADER
from data import formations, competences, experiences  # Importer les données depuis data.py
import datetime
import sys
import os

# Vérifiez que save_path est passé comme argument
if len(sys.argv) != 2:
    print("Usage: python template.py <save_path>")
    sys.exit(1)

filename = sys.argv[1]

# Créer le répertoire s'il n'existe pas
os.makedirs(os.path.dirname(filename), exist_ok=True)

doc = Document()

# Function to add a shading to a paragraph
def add_shading(paragraph, color="auto"):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    
# Fonction pour ajouter une bordure inférieure à un paragraphe
def add_bottom_border(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # Taille de la bordure
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '25fde9')  # Noir
    pBdr.append(bottom)
    pPr.append(pBdr)

# Function to add training entries with tab alignment
def add_training_entries(doc, trainings):
    for year, descriptions in trainings:
        descriptions_list = descriptions.split('\n')
        
        # Add the year and the first line of the description in one paragraph
        entry_paragraph = doc.add_paragraph()
        tab_stops = entry_paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(4), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
        entry_paragraph.paragraph_format.left_indent = Cm(0.8)
        
        year_run = entry_paragraph.add_run(year)
        year_run.font.name = 'Calibri'
        year_run.font.size = Pt(10)
        year_run.font.bold = True
        
        desc_run = entry_paragraph.add_run("\t" + descriptions_list[0])
        desc_run.font.name = 'Calibri'
        desc_run.font.size = Pt(10)
        
        # Add the remaining lines of the description as indented paragraphs
        for description in descriptions_list[1:]:
            desc_paragraph = doc.add_paragraph()
            desc_run = desc_paragraph.add_run("\t" + description)
            desc_run.font.name = 'Calibri'
            desc_run.font.size = Pt(10)
            desc_paragraph.paragraph_format.left_indent = Cm(4)
            
# Fonction pour ajouter des compétences et sous-compétences
def add_skills(doc, skills):
    # Ajouter un saut de ligne
    doc.add_paragraph()
    
    # Ajouter un titre pour les compétences
    skills_title_paragraph = doc.add_paragraph()
    skills_title_run = skills_title_paragraph.add_run('        COMPETENCES')
    skills_title_run.font.name = 'Calibri'
    skills_title_run.font.size = Pt(14)
    skills_title_run.font.bold = True
    skills_title_run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    skills_title_paragraph.alignment = 0  # Alignement à gauche
    skills_title_paragraph.paragraph_format.left_indent = -Cm(0.7)
    
    # Ajouter un ombrage au paragraphe de titre
    add_shading(skills_title_paragraph, color='9595A3')  # Ombrage violet

    # Ajouter chaque compétence et ses sous-compétences
    for skill, subskills in skills.items():
        # Ajouter le titre de la compétence
        skill_paragraph = doc.add_paragraph()
        skill_run = skill_paragraph.add_run(skill)
        skill_run.font.name = 'Calibri'
        skill_run.font.size = Pt(10)
        skill_run.font.bold = True
        skill_paragraph.paragraph_format.space_before = Pt(8)
        skill_paragraph.paragraph_format.space_after = Pt(2)
        
        # Ajouter les sous-compétences
        for subskill in subskills:
            subskill_paragraph = doc.add_paragraph(subskill, style='List Bullet')
            subskill_paragraph.paragraph_format.left_indent = Inches(0.5)
            subskill_run = subskill_paragraph.runs[0]
            subskill_run.font.name = 'Calibri'
            subskill_run.font.size = Pt(10)
            subskill_paragraph.paragraph_format.space_after = Pt(2)

# Fonction pour ajouter des expériences professionnelles
def add_experiences(doc, experiences):
    # Ajouter un saut de ligne
    doc.add_paragraph()
    
    # Ajouter un titre pour les expériences professionnelles
    exp_title_paragraph = doc.add_paragraph()
    exp_title_run = exp_title_paragraph.add_run('        EXPERIENCES PROFESSIONNELLES')
    exp_title_run.font.name = 'Calibri'
    exp_title_run.font.size = Pt(14)
    exp_title_run.font.bold = True
    exp_title_run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    exp_title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    exp_title_paragraph.paragraph_format.left_indent = -Cm(0.7)
    
    # Ajouter un ombrage au paragraphe de titre
    add_shading(exp_title_paragraph, color='9595A3')  # Ombrage violet

    # Ajouter chaque expérience professionnelle
    for experience in experiences:
        if len(experience) == 6:
            entreprise, periode, poste, description, technologies, contexte = experience
        else:
            entreprise, periode, poste, description, technologies = experience
            contexte = None
        
        # Ajouter un paragraphe pour l'entreprise et la durée sur la même ligne
        experience_paragraph = doc.add_paragraph()
        tab_stops = experience_paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        
        entreprise_run = experience_paragraph.add_run(entreprise.upper())
        entreprise_run.font.name = 'Gill Sans MT'
        entreprise_run.font.bold = True
        entreprise_run.font.color.rgb = RGBColor(0xac, 0xc0, 0xc6)
        entreprise_run.font.size = Pt(11)
        experience_paragraph.paragraph_format.left_indent = -Cm(0.25)
        experience_paragraph.add_run('\t')
        duree_run = experience_paragraph.add_run(periode)
        duree_run.font.name = 'Gill Sans MT'
        duree_run.font.bold = True
        duree_run.font.color.rgb = RGBColor(0xac, 0xc0, 0xc6)
        duree_run.font.size = Pt(11)
        
        # Ajouter une bordure inférieure à ce paragraphe
        add_bottom_border(experience_paragraph)
        
        # Réduire l'espacement après ce paragraphe
        experience_paragraph.paragraph_format.space_after = Pt(2)
        
        # Ajouter le poste en-dessous de l'entreprise et la durée, collé à la bordure
        poste_paragraph = doc.add_paragraph()
        poste_paragraph.paragraph_format.space_before = Pt(0)  # Réduire l'espace avant
        poste_paragraph.paragraph_format.space_after = Pt(6)   # Réduire l'espace après
        poste_run = poste_paragraph.add_run(poste.upper())
        poste_run.font.name = 'Cambria'
        poste_run.font.color.rgb = RGBColor(0x7d, 0x7d, 0x7d)
        poste_run.font.bold = True
        poste_run.font.size = Pt(11)
        poste_paragraph.paragraph_format.left_indent = -Cm(0.25)
        
        # Ajouter le contexte s'il existe
        if contexte:
            contexte_paragraph = doc.add_paragraph()
            contexte_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            contexte_run_bold = contexte_paragraph.add_run('Contexte : ')
            contexte_run_bold.font.name = 'Calibri'
            contexte_run_bold.font.size = Pt(10)
            contexte_run_bold.font.bold = True
            contexte_paragraph.paragraph_format.left_indent = -Cm(0.25)
            contexte_run = contexte_paragraph.add_run(contexte)
            contexte_run.font.name = 'Calibri'
            contexte_run.font.size = Pt(10)
        
        # Ajouter la section Responsabilité
        if description:
            responsabilite_paragraph = doc.add_paragraph('Responsabilité :')
            responsabilite_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            responsabilite_run = responsabilite_paragraph.runs[0]
            responsabilite_run.font.name = 'Calibri'
            responsabilite_run.font.size = Pt(10)
            responsabilite_run.font.bold = True
            responsabilite_paragraph.paragraph_format.left_indent = -Cm(0.25)
            responsabilite_paragraph.paragraph_format.space_after = Pt(2)
        
            # Ajouter la description
            for line in description:
                desc_paragraph = doc.add_paragraph(line, style='List Bullet')
                desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                desc_paragraph.paragraph_format.left_indent = Cm(1.6)
                desc_paragraph.paragraph_format.first_line_indent = -Cm(0.5)  # Indenter les lignes suivantes du texte
                desc_run = desc_paragraph.runs[0]
                desc_run.font.name = 'Calibri'
                desc_run.font.size = Pt(10)
        
        # Ajouter les technologies utilisées sur la même ligne
        if technologies:
            tech_paragraph = doc.add_paragraph()
            tech_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            tech_paragraph.paragraph_format.left_indent = -Cm(0.25)
            tech_paragraph.paragraph_format.space_before = Pt(0)
            tech_paragraph.paragraph_format.space_after = Pt(0)
            tech_run = tech_paragraph.add_run('Technologies utilisées : ')
            tech_run.font.name = 'Calibri'
            tech_run.font.size = Pt(10)
            tech_run.font.bold = True
            tech_run_bis = tech_paragraph.add_run(", ".join(technologies))
            tech_run_bis.font.name = 'Calibri'
            tech_run_bis.font.size = Pt(10)
            
        # Ajouter un saut de ligne entre les expériences
        doc.add_paragraph()

def main(save_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.21)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.29)
    section.right_margin = Cm(1.83)

    header_paragraph = doc.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    initials_run = header_paragraph.add_run('C.B')
    initials_run.font.name = 'Calibri'
    initials_run.font.size = Pt(18)
    initials_run.font.bold = True
    initials_run.font.color.rgb = RGBColor(255, 0, 0)
    header_paragraph.add_run('\t')
    poste_run = header_paragraph.add_run('TECHNICIEN')
    poste_run.font.name = 'Calibri'
    poste_run.font.size = Pt(18)
    poste_run.font.bold = True
    poste_run.font.color.rgb = RGBColor(255, 0, 0)

    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = header_paragraph.add_run()
    run.add_picture("antares_logo.png", width=Inches(1.5))

    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_paragraph.add_run()
    run.add_picture("bas_antares.png", width=Inches(6))

    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run('        FORMATIONS')
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    title_paragraph.alignment = 0
    title_paragraph.paragraph_format.left_indent = -Cm(0.7)
    add_shading(title_paragraph, color='9595A3')

    # Add training entries, skills, and experiences
    add_training_entries(doc, formations)
    add_skills(doc, competences)
    add_experiences(doc, experiences)

    # Save the document
    doc.save(filename)
    print(f"Document saved as {filename}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        save_path = sys.argv[1]
        main(save_path)
    else:
        print("Please provide a save path for the document.")
